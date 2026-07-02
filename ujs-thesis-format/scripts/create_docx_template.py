#!/usr/bin/env python3
"""Create a starter DOCX template for UJS thesis documents."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Cm, Pt


def set_run_font(
    run,
    east_asia: str,
    latin: str = "Times New Roman",
    size_pt: float | None = None,
    bold: bool = False,
    color: RGBColor | None = None,
):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_style_font(
    style,
    east_asia: str,
    latin: str = "Times New Roman",
    size_pt: float | None = None,
    bold: bool = False,
    color: RGBColor | None = None,
):
    style.font.name = latin
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size_pt is not None:
        style.font.size = Pt(size_pt)
    style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color


def add_bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_begin)

    run_instr = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run_instr._r.append(instr)

    run_sep = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(fld_sep)
    set_run_font(run_sep, "Times New Roman", size_pt=10)

    run_text = paragraph.add_run("1")
    set_run_font(run_text, "Times New Roman", size_pt=10)

    run_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.right_margin = Cm(2)
    section.different_first_page_header_footer = False

    header_p = section.header.paragraphs[0]
    header_p.text = ""
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_p.add_run("江苏大学本科毕业设计（论文）")
    set_run_font(run, "SimHei", size_pt=10.5, bold=True)
    add_bottom_border(header_p)

    footer_p = section.footer.paragraphs[0]
    footer_p.text = ""
    add_page_number(footer_p)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    set_style_font(normal, "FangSong", size_pt=12)
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = doc.styles[name]
        set_style_font(style, "SimHei", size_pt=size, bold=True, color=RGBColor(0, 0, 0))
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_starter_content(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run("毕业设计（论文）题目")
    set_run_font(run, "FZXiaoBiaoSong-B05S", size_pt=22, bold=True, color=RGBColor(0, 0, 0))

    fields = [
        "学院：",
        "专业班级：",
        "学生姓名：",
        "指导教师：",
        "指导教师职称：",
        "时间：",
    ]
    for field in fields:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        run = p.add_run(field)
        set_run_font(run, "FangSong", size_pt=12)

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("原创性声明", level=1)
    doc.add_paragraph("在此放置学校或学院要求的原创性声明正文。")
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("中文摘要", level=1)
    doc.add_paragraph("在此撰写不少于400字的中文摘要。")
    doc.add_paragraph("关键词：关键词一；关键词二；关键词三")
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph("Write the English abstract here.")
    doc.add_paragraph("Keywords: keyword one; keyword two; keyword three")
    doc.add_heading("目录", level=1)
    doc.add_paragraph("生成目录后保留至三级标题。")
    doc.add_heading("1 引言", level=1)
    doc.add_paragraph("在此撰写引言。")
    doc.add_heading("2 正文", level=1)
    doc.add_paragraph("在此撰写正文。")
    doc.add_heading("结论", level=1)
    doc.add_paragraph("在此撰写结论。")
    doc.add_heading("参考文献", level=1)
    for idx in range(1, 16):
        doc.add_paragraph(f"[{idx}] 作者. 题名[J]. 刊名, 年, 卷(期): 起止页码.")


def create_template(output: Path):
    doc = Document()
    configure_document(doc)
    configure_styles(doc)
    add_starter_content(doc)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description="Create a starter UJS thesis DOCX template.")
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    create_template(args.output)
    print(f"Created {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
